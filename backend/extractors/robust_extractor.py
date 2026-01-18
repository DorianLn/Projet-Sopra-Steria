"""
Extracteur robuste - optimisé pour multiples formats de CV
Basé sur cas réels : Sopra, étudiants, CV seniors
"""

import re
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
import pdfplumber
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ==========================================================
# EXTRACTION TEXTE BRUT
# ==========================================================

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(
                    x_tolerance=2,
                    y_tolerance=2
                )
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        logger.error(f"Erreur PDF: {e}")
        raise


def extract_text_from_docx(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
        text = ""

        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"

        return text

    except Exception as e:
        logger.error(f"Erreur DOCX: {e}")
        raise


def extract_text(file_path: str) -> str:
    path = Path(file_path)

    if path.suffix.lower() == ".pdf":
        return extract_text_from_pdf(file_path)

    if path.suffix.lower() == ".docx":
        return extract_text_from_docx(file_path)

    raise ValueError("Format non supporté")


# ==========================================================
# NETTOYAGE AVANCÉ
# ==========================================================

def clean_text(text: str) -> str:

    patterns = [
        r"C2\s*–\s*Usage restreint",
        r"C3\s*–\s*Confidentiel",
        r"<\?xml.*?\?>",
        r"<.*?>",
        r"©.*Sopra.*",
        r"\d+/\d+\s*–",
        r"DocumentFile.*",
        r"Page\s*\d+",
    ]

    for p in patterns:
        text = re.sub(p, "", text, flags=re.IGNORECASE)

    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\n\s*\n", "\n", text)

    return text.strip()


# ==========================================================
# DÉTECTION FLEXIBLE DES SECTIONS
# ==========================================================

SECTION_ALIASES = {
    "competences": [
        "COMPETENCES",
        "COMPÉTENCES",
        "COMPETENCES TECHNIQUES",
        "COMPETENCES FONCTIONNELLES",
        "SKILLS",
        "OUTILS"
    ],
    "formations": [
        "FORMATION",
        "FORMATIONS",
        "FORMATION - CERTIFICATION",
        "FORMATIONS ET LANGUES",
        "ETUDES",
        "EDUCATION"
    ],
    "experiences": [
        "EXPERIENCES",
        "EXPÉRIENCES",
        "EXPERIENCES PROFESSIONNELLES",
        "PARCOURS"
    ],
    "langues": [
        "LANGUES",
        "LANGUE(S)",
        "LANGUAGES"
    ],
    "loisirs": [
        "LOISIRS",
        "CENTRES D’INTÉRÊTS",
        "ACTIVITES EXTRA",
        "HOBBIES"
    ]
}


def normalize_title(title: str) -> str:
    title = title.upper().strip()

    if title.startswith("EXPERIENCE"):
        return "experiences"

    if title.startswith("SAVOIRS"):
        return "competences"

    for key, aliases in SECTION_ALIASES.items():
        for a in aliases:
            if a in title:
                return key

    return "contact"


def find_sections(text: str) -> Dict[str, str]:

    sections = {
        "contact": "",
        "competences": "",
        "experiences": "",
        "formations": "",
        "langues": "",
        "loisirs": ""
    }

    current = "contact"

    for line in text.split("\n"):
        clean = line.strip()

        if not clean:
            continue

        normalized = normalize_title(clean)

        # Si la ligne est un vrai titre connu
        if normalized in sections and clean.isalpha():
            current = normalized
            continue

        sections[current] += clean + "\n"

    return sections




# ==========================================================
# PARSING DES SECTIONS
# ==========================================================

def parse_contact(text: str) -> Dict[str, Any]:
    result = {
        "nom": None,
        "email": None,
        "telephone": None,
        "adresse": None,
        "titre_profil": None
    }

    email = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    if email:
        result["email"] = email.group(0)

    tel = re.search(r"(0|\+33)[\s\d\(\)\.]{8,20}", text)
    if tel:
        result["telephone"] = tel.group(0)

    # Détection d'adresse multi-lignes
    lines = text.split("\n")
    adresse_complete = None

    for i, line in enumerate(lines):
        if re.search(r"\d{5}\s+[A-Za-z\- ]+", line):
            adresse_complete = line.strip()

            # vérifier si la ligne précédente ressemble à une rue
            if i > 0 and re.search(r"\d+\s+\w+", lines[i-1]):
                adresse_complete = lines[i-1].strip() + " " + adresse_complete

            break

    if adresse_complete:
        result["adresse"] = adresse_complete

    for line in text.split("\n")[:10]:
        clean = line.strip()

        if any(m in clean.lower() for m in ["analyst", "développeur", "ingénieur", "consultant"]):
            continue

        if re.match(r"^[A-ZÉÈÊÀÂÎÔÙÛ][a-zéèêàâîôùû]+ [A-ZÉÈÊÀÂÎÔÙÛ]+", clean):
            result["nom"] = clean
            break

    
    # Si aucun nom détecté mais initiales présentes
    if result["nom"] is None:
        match = re.search(r"\b[A-Z]{2,4}\b", text)
        if match:
            result["nom"] = match.group(0)

    # NOUVEAU : détection titre profil
    lines = text.split("\n")

    for line in lines:
        low = line.lower()
        if any(k in low for k in ["étudiant", "developpeur", "ingénieur", "consultant", "stage"]):
            result["titre_profil"] = line.strip()
            break

    return result

# ==========================================================
# PARSING AVANCÉ DES SECTIONS
# ==========================================================
def split_by_real_sections(text: str) -> Dict[str, str]:

    sections = {
        "titre_profil": "",
        "nom": "",
        "competences_fonctionnelles": "",
        "competences_techniques": "",
        "experiences": "",
        "formations": "",
        "langues": ""
    }

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    current = None

    for i, line in enumerate(lines):
        lower = line.lower()

        if lower.startswith("langue"):
            current = "langues"

            contenu = re.sub(r"langue\(s\)", "", lower).strip()
            if contenu:
                sections[current] += contenu + "\n"

            continue

        if "compétences fonctionnelles" in lower:
            current = "competences_fonctionnelles"
            continue

        if "compétences techniques" in lower:
            current = "competences_techniques"
            continue

        if lower == "expériences":
            current = "experiences"
            continue

        if lower == "formation":
            current = "formations"
            continue

        # Si aucune section encore définie → on suppose contact
        if current is None:
            if not sections["titre_profil"]:
                sections["titre_profil"] = line
            elif not sections["nom"]:
                sections["nom"] = line
            continue

        sections[current] += line + "\n"

    return sections




# ==========================================================
# PARSING STRUCTURÉ SIMPLE (VERSION ANTÉRIEURE)
# ==========================================================
def parse_structured_cv(text: str) -> Dict[str, Any]:

    parts = split_by_real_sections(text)

    # Découpage titre_profil / nom quand ils sont concaténés
    titre = parts["titre_profil"]
    nom = parts["nom"]

    # Cas typique : "Business Analyst JLA"
    match = re.match(r"(.+)\s+([A-Z]{2,5})$", titre)

    if match:
        titre_profil = match.group(1).strip()
        nom = match.group(2).strip()
    else:
        titre_profil = titre

    # Récupération brute
    raw_experiences = [
        l.strip()
        for l in parts["experiences"].split("\n")
        if l.strip()
    ]

    raw_formations = [
        l.strip()
        for l in parts["formations"].split("\n")
        if l.strip()
    ]

    experiences = []
    formations = raw_formations.copy()

    for line in raw_experiences:

        # Formation si commence par "Formation 20xx"
        if re.match(r"^Formation\s+\d{4}", line, re.IGNORECASE):
            formations.append(line)

        # Certification ou diplôme explicite
        elif re.search(r"(SAFe|ISTQB|Certification|POEI|Dipl[oô]me)", line, re.IGNORECASE):
            formations.append(line)

        # Ligne qui commence par une année → formation
        elif re.match(r"^(19|20)\d{2}\s*[–-]", line):
            formations.append(line)

        else:
            experiences.append(line)


    data = {
        "contact": {
            "nom": nom,
            "email": None,
            "telephone": None,
            "adresse": None,
            "titre_profil": titre_profil
        },

        "competences": {
            "fonctionnelles": [
                l.strip()
                for l in parts["competences_fonctionnelles"].split("\n")
                if l.strip()
            ],
            "techniques": [
                l.strip()
                for l in parts["competences_techniques"].split("\n")
                if l.strip()
            ]
        },

        "experiences": experiences,

        "formations": formations,

        "langues": [
            l.strip()
            for l in parts["langues"].split("\n")
            if l.strip()
        ],

        "loisirs": [],

        "texte_brut": text
    }

    return data


def parse_competences(text: str) -> List[str]:

    competences = []
    current = ""

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            continue

        # Début d’une vraie compétence (ligne commençant par tiret ou mot-clé)
        if line.startswith("-") or ":" in line:
            if current:
                competences.append(current.strip())
            current = line.strip("- ").strip()
        else:
            # continuation de la ligne précédente
            current += " " + line

    if current:
        competences.append(current.strip())

    return competences

def split_competences(competences: List[str]) -> Dict[str, List[str]]:

    techniques = []
    fonctionnelles = []

    is_soft = False

    for comp in competences:

        # Si la ligne contient Savoirs-être, on la nettoie et on bascule APRÈS
        if "Savoirs-être" in comp:
            comp = comp.replace("Savoirs-être", "").strip()
            if comp:
                techniques.append(comp)
            is_soft = True
            continue

        if is_soft:
            fonctionnelles.append(comp)
        else:
            techniques.append(comp)

    return {
        "techniques": techniques,
        "fonctionnelles": fonctionnelles
    }


def parse_formations(text: str) -> List[str]:

    formations = []

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            continue

        formations.append(line)

    return formations



def parse_experiences(text: str) -> List[str]:

    experiences = []
    current = ""

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            continue

        # Si la ligne contient une année → début d’une nouvelle expérience
        if re.search(r"(19|20)\d{2}", line):

            if current:
                experiences.append(current.strip())

            current = line
            continue

        # Si la ligne commence par un tiret → description liée à l’expérience en cours
        if line.startswith("-"):
            current += " " + line.strip("- ").strip()
            continue

    # Ajouter la dernière expérience
    if current:
        experiences.append(current.strip())

    return experiences


def parse_langues(text: str):

    result = []

    for line in text.split("\n"):
        line = line.strip()

        if "français" in line.lower() or "anglais" in line.lower():
            # On enlève les tirets ou puces éventuelles
            clean = line.strip("-• ").strip()
            result.append(clean)

    return result

def parse_loisirs(text: str) -> List[str]:
    loisirs = []

    for line in text.split("\n"):
        line = line.strip("-• ")
        if line:
            loisirs.append(line)

    return loisirs


# ==========================================================
# PIPELINE PRINCIPAL
# ==========================================================

def extract_cv_robust(file_path: str) -> Dict[str, Any]:

    logger.info(f"Extraction de: {file_path}")

    raw = extract_text(file_path)
    clean = clean_text(raw)

    # Détection d'un CV déjà bien structuré type JLA
    if re.search(r"Compétences\s+fonctionnelles", clean, re.IGNORECASE) \
    and re.search(r"Compétences\s+techniques", clean, re.IGNORECASE):

        logger.info("CV structuré détecté (type JLA)")
        return parse_structured_cv(clean)


    # Sinon on garde TON ANCIEN PARSER
    sections = find_sections(clean)
    competences_raw = parse_competences(sections.get("competences", ""))

    data = {
        "contact": parse_contact(sections.get("contact", "")),
        "competences": split_competences(competences_raw),
        "formations": parse_formations(sections.get("formations", "")),
        "experiences": parse_experiences(sections.get("experiences", "")),
        "langues": parse_langues(sections.get("langues", "")),
        "loisirs": parse_loisirs(sections.get("loisirs", "")),
        "texte_brut": clean
    }

    return data


if __name__ == "__main__":
    import json
    res = extract_cv_robust("test.pdf")
    print(json.dumps(res, indent=2, ensure_ascii=False))
