import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from typing import Dict, List, Optional, Any
import re


# ---------------------------
#      HELPERS VISUELS
# ---------------------------

def add_horizontal_line(paragraph, color="7030A0"):
    """Ajoute une ligne horizontale style Sopra (sous le nom du profil)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    pbdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color)  # violet Sopra par défaut
    pbdr.append(bottom)

    pPr.append(pbdr)
# ---------------------------
#      DOCX FORMATTING      
# ---------------------------
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, space_after=6, line_spacing=1.15):
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    # Espacement après paragraphe
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(space_after * 20))  # valeur en twips
    spacing.set(qn("w:line"), str(int(line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")

    pPr.append(spacing)

def set_cell_shading(cell, color):
    """Définit la couleur de fond d'une cellule."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# ---------------------------
#   CLASSIFICATION COMPÉTENCES
# ---------------------------

TECH_KEYWORDS = [
    "python","java","c++","c#","javascript","typescript","sql","docker","kubernetes",
    "aws","azure","gcp","cloud","devops","ci","cd","ci/cd","git","react","angular",
    "vue","node","node.js","django","flask","fastapi","spring","github","gitlab",
    "linux","mongodb","postgresql","mysql","redis","kafka","api","rest","graphql",
    ".net","php","ruby","go","rust","swift","kotlin","scala","html","css","sass"
]

FONC_KEYWORDS = [
    "communication","gestion","travail","team","leadership","équipe","equipe",
    "organisation","planning","collaboration","management","autonomie","analyse",
    "résolution","problèmes","relation","client","négociation","presentation"
]

def classify_competences(lst):
    if not isinstance(lst, list):
        return [], []

    comp_fonct, comp_tech = [], []
    seen_fonct, seen_tech = set(), set()

    for c in lst:
        if not isinstance(c, str):
            continue
        low = c.lower().strip()
        
        if not low or len(low) < 2:
            continue

        if any(k in low for k in TECH_KEYWORDS):
            if low not in seen_tech:
                seen_tech.add(low)
                comp_tech.append(c)
        elif any(k in low for k in FONC_KEYWORDS):
            if low not in seen_fonct:
                seen_fonct.add(low)
                comp_fonct.append(c)
        else:
            if low not in seen_tech:
                seen_tech.add(low)
                comp_tech.append(c)

    return comp_fonct, comp_tech

def bullets(lst, fallback="Non renseigné"):
    """Génère une liste à puces avec fallback intelligent"""
    if not lst:
        return f"• {fallback}"
    clean_items = [str(v).strip() for v in lst if v and str(v).strip()]
    if not clean_items:
        return f"• {fallback}"
    return "\n".join(f"• {v}" for v in clean_items)

def format_experiences(exps):

    if not isinstance(exps, list) or not exps:
        return "• Aucune expérience renseignée"

    result = []

    for exp in exps:
        exp = exp.strip()
        if not exp:
            continue

        # ----- CAS SPÉCIAL : expériences contenant des projets -----
        if "Projets" in exp:
            parts = exp.split("Projets", 1)

            header = parts[0].strip().rstrip(" :")
            details = parts[1].strip()

            # Titre principal
            result.append(f"{header} : Projets")
            result.append("")

            # Nettoyage du mot "Projets" résiduel
            details = details.replace("Projets", "").strip()

            # Découper chaque projet à partir de "Projet individuel"
            projets = re.split(r"(Projet individuel)", details)

            for i in range(1, len(projets), 2):
                projet = projets[i] + projets[i + 1]

                # Nettoyage COMPLET des caractères invisibles
                projet = projet.replace("\n", " ")
                projet = projet.replace("\t", " ")
                projet = projet.replace("\r", " ")
                projet = projet.replace("\xa0", " ")   # espace insécable

                # Supprimer tous espaces multiples
                projet = re.sub(r"[ ]{2,}", " ", projet)

                # Nettoyage final
                projet = projet.strip()

                result.append("• " + projet)

            result.append("")
            continue

        # ----- CAS NORMAL (autres expériences) -----
        split_keywords = ["Programme", "Stage"]

        split_index = -1
        for kw in split_keywords:
            idx = exp.find(kw)
            if idx != -1:
                split_index = idx
                break

        if split_index != -1:
            header = exp[:split_index].strip()
            details = exp[split_index:].strip()

            result.append(header)
            result.append(f"• {details}")
        else:
            result.append(f"• {exp}")

        result.append("")

    return "\n".join(result).strip()




def format_formations(forms):

    if not forms:
        return "• Aucune formation renseignée"

    sorted_forms = sort_formations_by_date(forms)

    lines = []

    for f in sorted_forms:

        # Cas 1 : formation en string (nouveau format)
        if isinstance(f, str):
            lines.append(f"• {f}")
            continue

        # Cas 2 : ancien format dict
        if isinstance(f, dict):
            etablissement = f.get("etablissement", "Établissement non précisé")
            dates = normalize_date_display(f.get("dates"))
            diplome = f.get("diplome")

            if diplome:
                line = f"• {diplome} – {etablissement} ({dates})"
            else:
                line = f"• {etablissement} ({dates})"

            lines.append(line)

    return "\n".join(lines) if lines else "• Aucune formation renseignée"



def sort_experiences_by_date(exps):
    """Trie les expériences par date (anti-chronologique) en acceptant string ou dict."""

    def extract_year(exp):

        if isinstance(exp, str):
            match = re.search(r"(19|20)\d{2}", exp)
            if match:
                return int(match.group(0))
            return 0

        if isinstance(exp, dict):
            dates = str(exp.get("dates", ""))
            match = re.search(r"(19|20)\d{2}", dates)
            if match:
                return int(match.group(0))
            return 0

        return 0

    return sorted(exps, key=extract_year, reverse=True)


def sort_formations_by_date(forms):
    """Trie les formations par date (anti-chronologique) en acceptant string ou dict."""

    def extract_year(form):

        # Cas 1 : la formation est une STRING
        if isinstance(form, str):
            match = re.search(r"(19|20)\d{2}", form)
            if match:
                return int(match.group(0))
            return 0

        # Cas 2 : la formation est un DICT (ancien format)
        if isinstance(form, dict):
            dates = str(form.get("dates", ""))
            match = re.search(r"(19|20)\d{2}", dates)
            if match:
                return int(match.group(0))
            return 0

        return 0

    return sorted(forms, key=extract_year, reverse=True)



def normalize_date_display(date_str: Optional[str]) -> str:
    """Normalise l'affichage des dates."""
    if not date_str:
        return "Dates non précisées"
    
    date_str = str(date_str).strip()
    
    # Normaliser les tirets
    date_str = re.sub(r'\s*[-–—]\s*', ' – ', date_str)
    
    # Normaliser "Présent"
    date_str = re.sub(r'(?i)\b(actuel|actuellement|aujourd\'?hui)\b', 'Présent', date_str)
    
    return date_str


def format_contact(contact):
    """Formate les informations de contact"""
    if not isinstance(contact, dict):
        return ""
    
    lines = []
    if contact.get("email"):
        lines.append(f"Email : {contact['email']}")
    if contact.get("telephone"):
        lines.append(f"Tél : {contact['telephone']}")
    if contact.get("adresse"):
        lines.append(f"Adresse : {contact['adresse']}")
    
    return "\n".join(lines) if lines else "Contact non renseigné"

# ---------------------------
#      DOCX HEADER
# ---------------------------
def add_header_to_document(doc, titre_profil, nom):
    """
    Ajoute le titre de profil et le nom dans l'en-tête du document
    à partir de la deuxième page.
    """
    for section in doc.sections:
        # Activer un header différent pour la première page
        section.different_first_page_header_footer = True
        
        # Marges pour l'en-tête
        section.top_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        # En-tête par défaut (pages 2+)
        header = section.header
        
        # Nettoyer l'en-tête existant
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # Ajouter le titre profil (sans saut de ligne avant)
        p_titre = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run_titre = p_titre.add_run(titre_profil)
        run_titre.bold = True
        run_titre.font.size = Pt(12)
        run_titre.font.color.rgb = RGBColor(0, 0, 0) 
        p_titre.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p_titre, space_after=2, line_spacing=1.0)
        
        # Ajouter le nom
        p_nom = header.add_paragraph()
        run_nom = p_nom.add_run(nom)
        run_nom.bold = True
        run_nom.font.size = Pt(10)
        run_nom.font.color.rgb = RGBColor(0, 0, 0)
        p_nom.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p_nom, space_after=6, line_spacing=1.0)
        
        # Ligne de séparation orange Sopra
        p_line = header.add_paragraph()
        add_horizontal_line(p_line, color="F08650")
# ---------------------------
#   GENERATE DOCX FINAL
# ---------------------------
def generate_sopra_docx(cv_data, output_path):
    template_path = "templates/sopra_template.docx"

    if not os.path.exists(template_path):
        raise FileNotFoundError("Template DOCX introuvable")

    doc = Document(template_path)

    contact = cv_data.get("contact", {}) or {}

    titre_profil = contact.get("titre_profil") or "Profil Collaborateur"
    nom = contact.get("nom") or "Nom Prénom"

    competences = cv_data.get("competences", {}) or {}

    comp_tech = competences.get("techniques", [])
    comp_fonct = competences.get("fonctionnelles", [])

    mapping = {
        "{{TITRE_PROFIL}}": titre_profil,
        "{{NOM_PRENOM}}": nom,
        "{{COMP_FONCT}}": bullets(comp_fonct, "Aucune compétence fonctionnelle"),
        "{{COMP_TECH}}": bullets(comp_tech, "Aucune compétence technique"),
        "{{EXPERIENCES}}": format_experiences(cv_data.get("experiences", [])),
        "{{FORMATIONS_CERTIFICATIONS}}": format_formations(cv_data.get("formations", [])),
        "{{LANGUES}}": bullets(cv_data.get("langues", []), "Non renseigné"),
    }

    # --------- REMPLACEMENT CONTENU ---------
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:

                p.clear()

                # ----- CAS SPÉCIAL EXPERIENCES -----
                if key == "{{EXPERIENCES}}":

                    lines = val.split("\n")
                    # ---- AJOUT D'UN ESPACE AVANT LA PREMIÈRE EXPÉRIENCE ----
                    p.insert_paragraph_before("")

                    for line in lines:

                        if not line.strip():
                            doc.add_paragraph("")
                            continue

                        new_p = p.insert_paragraph_before("")
                        run = new_p.add_run(line)

                        # TITRE → aligné à gauche
                        if not line.strip().startswith("•"):
                            new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run.bold = True

                        # DÉTAIL → justifié
                        else:
                            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            new_p.paragraph_format.left_indent = Cm(1.0)

                        set_paragraph_spacing(new_p, space_after=6, line_spacing=1.15)

                    # on vide le paragraphe template
                    p.text = ""

                # ----- CAS NORMAL -----
                else:
                    run = p.add_run(val if val else "Non renseigné")

                    if key == "{{TITRE_PROFIL}}":
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(20)
                            run.font.color.rgb = RGBColor(77, 27, 130)
                        set_paragraph_spacing(p, space_after=4, line_spacing=1.1)

                    elif key == "{{NOM_PRENOM}}":
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(16)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        set_paragraph_spacing(p, space_after=10, line_spacing=1.1)

                    else:
                        set_paragraph_spacing(p, space_after=8, line_spacing=1.2)

    # Traiter les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in mapping.items():
                        if key in p.text:
                            p.text = p.text.replace(key, val if val else "Non renseigné")
                    set_paragraph_spacing(p, space_after=8, line_spacing=1.2)

    # ✅ AJOUTER L'EN-TÊTE POUR LES PAGES 2+
    add_header_to_document(doc, titre_profil, nom)

    doc.save(output_path)
    return output_path


def generate_cv_documents(cv_data: Dict, output_dir: str, filename: str = None) -> Dict[str, str]:
    """
    Génère les documents CV (DOCX et PDF) à partir des données structurées.
    
    Args:
        cv_data: Données CV structurées (contact, formations, experiences, etc.)
        output_dir: Dossier de sortie
        filename: Nom du fichier (sans extension)
    
    Returns:
        Dict avec les chemins des fichiers générés: {"docx": path, "pdf": path}
    """
    from datetime import datetime
    
    # Créer le dossier si nécessaire
    os.makedirs(output_dir, exist_ok=True)
    
    # Générer le nom de fichier
    if not filename:
        contact = cv_data.get("contact", {}) or {}
        nom = contact.get("nom", "CV")
        # Nettoyer le nom pour le fichier
        nom_clean = re.sub(r'[^\w\s-]', '', nom).strip().replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"CV_{nom_clean}_{timestamp}"
    
    docx_path = os.path.join(output_dir, f"{filename}.docx")
    pdf_path = os.path.join(output_dir, f"{filename}.pdf")
    
    result = {"docx": None, "pdf": None}
    
    # Générer le DOCX
    try:
        generate_sopra_docx(cv_data, docx_path)
        result["docx"] = docx_path
        print(f"✓ DOCX généré: {docx_path}")
    except Exception as e:
        print(f"❌ Erreur DOCX: {e}")
        raise
    
    # Générer le PDF
    try:
        from generators.docx_to_pdf import convert_docx_to_pdf
        convert_docx_to_pdf(docx_path, pdf_path)
        result["pdf"] = pdf_path
        print(f"✓ PDF généré: {pdf_path}")
    except ImportError:
        try:
            from docx_to_pdf import convert_docx_to_pdf
            convert_docx_to_pdf(docx_path, pdf_path)
            result["pdf"] = pdf_path
            print(f"✓ PDF généré: {pdf_path}")
        except Exception as e:
            print(f"⚠️ PDF non généré (docx2pdf indisponible): {e}")
    except Exception as e:
        print(f"⚠️ Erreur PDF: {e}")
    
    return result


def validate_cv_data(cv_data: Dict) -> List[str]:
    """
    Valide les données CV avant génération.
    
    Returns:
        Liste des avertissements/erreurs
    """
    warnings = []
    
    # Vérifier les champs obligatoires
    contact = cv_data.get("contact", {})
    if not contact:
        warnings.append("Contact manquant")
    else:
        if not contact.get("nom"):
            warnings.append("Nom manquant dans contact")
        if not contact.get("email"):
            warnings.append("Email manquant dans contact")
    
    # Vérifier les formations
    formations = cv_data.get("formations", [])
    if not formations:
        warnings.append("Aucune formation")
    else:
        for i, f in enumerate(formations):
            if isinstance(f, dict):
                if not f.get("etablissement") and not f.get("diplome"):
                    warnings.append(f"Formation {i+1}: établissement et diplôme manquants")
    
    # Vérifier les expériences
    experiences = cv_data.get("experiences", [])
    if not experiences:
        warnings.append("Aucune expérience")

    else:
        for i, e in enumerate(experiences):
            if not e.get("entreprise"):
                warnings.append(f"Expérience {i+1}: entreprise manquante")
    
    return warnings


def preview_cv_content(cv_data: Dict) -> str:

    lines = []

    contact = cv_data.get("contact", {})
    lines.append(f"=== {contact.get('nom', 'N/A')} ===")
    lines.append(f"Email: {contact.get('email', 'N/A')}")
    lines.append(f"Tél: {contact.get('telephone', 'N/A')}")
    lines.append("")

    lines.append("--- FORMATIONS ---")
    for f in cv_data.get("formations", [])[:3]:
        lines.append(f"  • {f}")

    lines.append("")
    lines.append("--- EXPÉRIENCES ---")
    for e in cv_data.get("experiences", [])[:3]:
        lines.append(f"  • {e}")

    lines.append("")
    lines.append("--- COMPÉTENCES ---")

    competences = cv_data.get("competences", {})

    tech = competences.get("techniques", [])
    fonct = competences.get("fonctionnelles", [])

    lines.append("Techniques : " + ", ".join(tech))
    lines.append("Fonctionnelles : " + ", ".join(fonct))

    return "\n".join(lines)
